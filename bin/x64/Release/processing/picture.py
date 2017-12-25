# -*- coding: utf-8 -*-
"""
Created on Thu Dec  7 09:24:03 2017

@author: Frank
"""

#from docx import Document
#from docx.shared import Pt
from docx.shared import Inches
#from docx.shared import RGBColor
#from docx.oxml.ns import qn
#from docx.text import paragraph
import docx
import os
#from misc import switch
#from lxml import etree
from config import db,POSITION,CERTIFICATE_SIZE,CERTIFICATE
'''
import cv2

def resave(inpath,outpath=None):
    img = cv2.imread(inpath)
    if outpath is None:
        outpath = inpath
    cv2.imwrite(outpath,img)
'''

from PIL import Image

def resave(inpath,outpath=None):
    img = Image.open(inpath)
    if outpath is None:
        outpath = inpath
    img.save(outpath)
    
def remove_space(text,mode='all'):
    if mode == 'all':
        return ''.join(text.split())#text.replace(" ", "")
    elif mode == 'ending or leading':
        return text.strip()

def delet_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

from xpinyin import Pinyin
p = Pinyin()

def chinese2english(chinese):
    return p.get_pinyin(chinese)[0].upper()

import glob

def mk_img_list(section_name,positions=[]):
    positions = positions
    categories = []
    if '身份证明' in  section_name:
        positions = ['法人']
        categories = ['身份证']
    elif '委托书' in section_name:
        if '（二）' in section_name:
            positions = ['委托人1','委托人2']
        else:
            positions = ['法人','委托人1']
        categories = ['身份证']
    elif '组织机构' in section_name:
        positions = list(db['project_members'].keys())
        positions = [position for position in positions if not position in ['法人','联系人','委托人']]
    elif '简历' in section_name:
        categories = CERTIFICATE
    #print(positions)
    names = [db['project_members'][position] for position in positions]
    #print(names)
    imgs = []
    for name in names:
        #\\W
        subdir_name = chinese2english(name)
        subpath = os.path.join(db['人员证件路径'],subdir_name)
        #print(subpath)
        if not os.path.exists(subpath):
            return imgs
        #\\W\\旺旺
        subpath_ = ''
        for subdir_ in os.listdir(subpath):
            if name in subdir_:
                subpath_ =  os.path.join(subpath,subdir_)
        #print(subpath_)
        if subpath_ is '':
            return imgs
        #\\W\\旺旺\\
        for category in categories:
            paths = glob.iglob(subpath_+'\\**\\' + '*'+category+'*.j*pg', recursive=True)
            paths = list(paths)
            #print(paths)
            if len(paths) == 0:
                continue
            for fullpath in paths:
                imgs.append([fullpath,category])
    return imgs

def set_run(run):
    #run.font.name = u'微软雅黑'
    #run.font.size = 10.5
    #run.font.underline = True
    #run.font.bold = True
    #run.font.italic = True
    #run.font.underline = WD_UNDERLINE.DOT_DASH
    #run.font.color.rgb =  RGBColor(0x42,0x24,0xE9)
    pass

def set_cur_member(text):
    for position in POSITION:
        if position in text:
            return position
    return None
 
from docx.enum.text import WD_BREAK

def insert_pictures(section_name,filepath):
    if not os.path.exists(filepath):
        print('such a file not exists')
        return
    doc =  docx.Document(filepath)
    #cur_member = ''
    i = 0
    #for i,para in enumerate(doc.paragraphs):
    while i < len(doc.paragraphs)-1:
        #print(i)
        para = doc.paragraphs[i]
        text = remove_space(para.text)
        #key =  set_cur_member(text)
        #if key is not None:
        #    cur_member = key
        if any(key in section_name for key in ['法定代表人身份证明','法定代表人','身份证明','授权委托书','授权委托书（一）格式','授权委托书（二）格式']): 
            #print(para.text)
            if any(key in text for key in ['特此证明','特此委托']):
                #p_e = paragraph._element
                next_p = doc.paragraphs[i+1]
                while len(next_p.text) == 0:
                    delet_paragraph(next_p)
                    next_p = doc.paragraphs[i+1]
                #set title
                #previous_p = next_p.insert_paragraph_before()
                #run = previous_p.add_run(u'法定代表人身份证证明：')
                #set_run(run)
                #
                previous_p = next_p.insert_paragraph_before()
                for img_path,category in mk_img_list(section_name):
                    run = previous_p.add_run()
                    run.add_picture(img_path,width=Inches(CERTIFICATE_SIZE[category]))
            elif '复印件' in text and '附' not in text:
                next_p = doc.paragraphs[i+1]
                while len(next_p.text) == 0:
                    delet_paragraph(next_p)
                    next_p = doc.paragraphs[i+1]
                previous_p = next_p.insert_paragraph_before()
                for img_path,category in mk_img_list(section_name):
                    run = previous_p.add_run()
                    run.add_picture(img_path,width=Inches(CERTIFICATE_SIZE[category]))
        if '简历' in section_name:
            if '附：' in text and '证件' in text:
                # add page-break
                page_break_p = para.insert_paragraph_before()
                page_break_p.add_run().add_break(WD_BREAK.PAGE)
                i += 1
                #
                next_p = doc.paragraphs[i+1]
                position = set_cur_member(remove_space(para.text))
                img_path_list = mk_img_list(section_name,[position])
                para.text = position + '证件'
                if len(img_path_list) > 0:
                    #print(position)
                    category_prev = '无'
                    for img_path,category in img_path_list:
                        #print(img_path,img_size)
                        # add text
                        if category_prev is not category:
                            title_p = next_p.insert_paragraph_before()
                            title_p.add_run(position+category)
                            i += 1
                        # sometimes it's stuck for images with height longer than width
                        # it is better to rotate such an image.
                        # add picture
                        if category_prev is not category:
                            picture_p = next_p.insert_paragraph_before()
                            i += 1
                        resave(img_path)
                        picture_p.add_run().add_picture(img_path,width=Inches(CERTIFICATE_SIZE[category]))
                        #run = previous_p.add_run()
                        #run.add_picture(img_path,width=Inches(CERTIFICATE_SIZE[category]))
                        
                        # add page-break
                        if category_prev is not category:
                            page_break_p = next_p.insert_paragraph_before()
                            page_break_p.add_run().add_break(WD_BREAK.PAGE)
                            i += 1          
                        #
                        category_prev = category
                    #print('adding end')
                    i += 1
        i += 1
    #print('saving....')            
    doc.save(filepath)